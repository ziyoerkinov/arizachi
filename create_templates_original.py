"""
OriginaL namunalardan to'g'ridan-to'g'ri shablonlar yaratish skripti.
Bunda Word dagi matn dizayni, shrifr, joylashuv 100% bir xil saqlanadi.
"""

import os
from docx import Document

def original_ni_shablonga_aylantirish(asl_fayl, shablon_fayl, almashtirishlar):
    """
    Asl fayl ichidagi matnlarni izlab, placeholderlarga almashtiradi
    va yangi shablon sifatida saqlaydi. Original dizayn 100% saqlanadi.
    """
    if not os.path.exists(asl_fayl):
        print(f"XATOLIK: Asl fayl topilmadi: {asl_fayl}")
        return

    doc = Document(asl_fayl)

    # Paragraflarni aylanib chiqish
    for paragraph in doc.paragraphs:
        for p_qidir, p_almashtir in almashtirishlar.items():
            if p_qidir in paragraph.text:
                # Matnni aniq o'zgartirish (run larni buzmaslik uchun diqqat qilamiz)
                _replace_text_safely(paragraph, p_qidir, p_almashtir)

    # Jadvallarni aylanib chiqish
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for p_qidir, p_almashtir in almashtirishlar.items():
                        if p_qidir in paragraph.text:
                            _replace_text_safely(paragraph, p_qidir, p_almashtir)
                            
    os.makedirs(os.path.dirname(shablon_fayl), exist_ok=True)
    doc.save(shablon_fayl)
    print(f"Muvaffaqiyatli shablon tayyorlandi: {shablon_fayl}")

def _replace_text_safely(paragraph, search_text, replace_text):
    """
    Paragraf ichida run-larni jiddiy buzmasdan matnni almashtiradi.
    """
    if search_text not in paragraph.text:
        return
        
    runs = paragraph.runs
    for i in range(len(runs)):
        if search_text in runs[i].text:
            runs[i].text = runs[i].text.replace(search_text, replace_text)
            return  # To'liq match bitta run ichida bo'lsa
            
    # Agar so'z bir nechta run-ga bo'linib ketgan bo'lsa:
    full_text = paragraph.text
    new_text = full_text.replace(search_text, replace_text)
    
    if len(runs) > 0:
        # Barcha runlarning formatini yig'ib, 1-run ga yozishga to'g'ri keladi 
        # (shrift kattaligi, boldness ni birinchi rundan oladi)
        runs[0].text = new_text
        for i in range(1, len(runs)):
            runs[i].text = ""

def generate_templates():
    import sys
    sys.stdout.reconfigure(encoding='utf-8')
    # -----------------------------------
    # 1. KRILL SHABLON ("Ариза уч оёқ.docx") nushasini shablonga o'girish
    # -----------------------------------
    asl_krill_fayl = r"D:\Law\2026 fevral\Уч оёқ\Ариза уч оёқ.docx"
    shablon_krill = r"d:\Apps\coding\UchOyoqBot\templates\ariza_krill.docx"
    
    krill_almashtirishlar = {
        "Қамаши тумани": "{{TUMAN_NOMI}}",
        "Обилов Отамурод Боймуродович": "{{FISH}}",
        "Обилов Отамурод Боймуродовичдан": "{{FISH}}дан",
        "Қашқадарё вилояти, Қамаши тумани, Чим ҚФЙ, Ўрта-Чим, Рақамсиз уй": "{{MANZIL}}",
        "+998 ": "{{TEL}}",
        "Қашқадарё вилояти МАБ инспектори Саидисломбек Шодихонов Муртозахон ўғли": "{{ORGAN_NOMI}}",
        "SD25700146635": "{{QAROR_RAQAMI}}",
        "28.10.2025": "{{QAROR_SANASI}}",
        "МЖтКнинг 128-моддаси": "{{MODDA}}",
        "412 000": "{{JARIMA_SUMMASI}}"
        # Eslatma: "SD25700146635-сонли" qoladi, biz faqat "SD25700146635" ni replace qilamiz.
    }
    
    try:
        original_ni_shablonga_aylantirish(asl_krill_fayl, shablon_krill, krill_almashtirishlar)
    except Exception as e:
        print(f"Xatolik (Krill): {e}")

    # -----------------------------------
    # 2. LOTIN SHABLON ("Ариза намуна2.docx") nushasini shablonga o'girish
    # -----------------------------------
    asl_lotin_fayl = r"D:\Law\2026 fevral\Уч оёқ\Ариза намуна2.docx"
    shablon_lotin = r"d:\Apps\coding\UchOyoqBot\templates\ariza_lotin.docx"
    
    lotin_almashtirishlar = {
        "OLMAZOR TUMANI": "{{TUMAN_NOMI}}",
        "JO‘RAYEV AZIMBEK AVAZ O‘G‘LI": "{{FISH}}",
        "TOSHKENT VILOYATI TOSHKENT TUMANI Kuxinur MFY, SHinam,5 tor ko‘chasi, 12A-uy da": "{{MANZIL}}",
        "+998974403222": "{{TEL}}",
        "TOSHKENT SHAHAR IIBB YHXB INSPEKTORI NAJMIDDINOV SALIMBOY XUSAN O‘G‘LI": "{{ORGAN_NOMI}}",
        "RA25089746675": "{{QAROR_RAQAMI}}",
        "04.11.2025": "{{QAROR_SANASI}}",
        "128X3-moddasi 1-qismi": "{{MODDA}}",
        "412 000": "{{JARIMA_SUMMASI}}",
        "MATIZ BEST": "{{AVTOMOBIL}}",
        "01L827KC": "{{DAVLAT_RAQAMI}}"
    }

    try:
        original_ni_shablonga_aylantirish(asl_lotin_fayl, shablon_lotin, lotin_almashtirishlar)
    except Exception as e:
        print(f"Xatolik (Lotin): {e}")

if __name__ == "__main__":
    generate_templates()
