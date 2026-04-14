"""
Shablon yaratish skripti — mavjud ariza faylini asos qilib, 
placeholder qiymatlar bilan yangi shablon yaratadi.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_template():
    """Krill alifbosidagi ariza shablonini yaratish."""
    doc = Document()
    
    # Sahifa sozlamalari
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    
    # Default shrift
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    # ===== HEADER =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Жиноят ишлари бўйича {{TUMAN_NOMI}} судига')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('{{FISH}}дан')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Манзил: {{MANZIL}}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Тел: {{TEL}}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    # ===== TITLE =====
    doc.add_paragraph()  # Bo'sh qator
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('А Р И З А')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('(ЙҲХБ қарорини бекор қилиб, иш юритишни тугатиш тўғрисида)')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_paragraph()  # Bo'sh qator
    
    # ===== BODY =====
    body_paragraphs = [
        '\t\t{{ORGAN_NOMI}}нинг {{QAROR_SANASI}} йилдаги {{QAROR_RAQAMI}}-сонли жарима солиш тўғрисидаги қарорига кўра махсус автоматлаштирилган фото ва видео қайд этиш техника воситалари орқали қайд этилган қоида бузилиши кўриб чиқилиб, МЖтКнинг {{MODDA}}сига асосан мени айбдор деб топиб {{JARIMA_SUMMASI}} сўм миқдорида жарима қўлланилган. Ушбу ҳуқуқбузарлик махсус автоматлаштирилган фото ва видео қайд этиш тизими орқали аниқланганлиги баён этилган.',
        
        'Ушбу жарима солиш тўғрисидаги қарор Ўзбекистон Республикасининг амалдаги қонунчилик талаблари қўпол равишда бузилган ҳолда расмийлаштирилган ва қуйидаги асосларга кўра бекор қилиниши лозим:',
        
        'Биринчидан, Ўзбекистон Республикасининг Маъмурий жавобгарлик тўғрисидаги кодекси 309¹-моддасига кўра фақатгина махсус автоматлаштирилган фото ва видео қайд этиш техника воситалари орқали қайд этилган йўл ҳаракати қоидалари бузилганлиги бўйича жарима солиш тўғрисидаги қарор маъмурий ҳуқуқбузарлик содир этган шахснинг иштирокисиз чиқарилиши белгиланган бўлиб, бунда кўчма фоторадар ва мобиль комплекслар орқали аниқланган қоидабузарликлар назарда тутилмаган.',
        
        'Жарима солиш тўғрисидаги қарор махсус автоматлаштирилган фото ва видео қайд этиш техника воситалари қўлланилган ҳолда олинган материаллар илова қилиниб, юридик кучга эга бўлган электрон ҳужжат шаклида расмийлаштирилади ҳамда ушбу ҳужжатни тузган ваколатли шахснинг электрон рақамли имзоси билан тасдиқланади.',
        
        'Шунингдек, Вазирлар Маҳкамасининг 2018 йил 1 декабрдаги 975-сон қарорига мувофиқ қабул қилинган низомнинг 37-бандида ҳам «Жарима солиш тўғрисидаги қарор махсус автоматлаштирилган фото ва видео қайд этиш техника воситалари қўлланилган ҳолда олинган материаллар илова қилиниб, юридик кучга эга бўлган электрон ҳужжат шаклида расмийлаштирилади ҳамда ушбу ҳужжатни тузган ваколатли шахснинг электрон рақамли имзоси билан тасдиқланади» деб белгиланган.',
        
        'Юқоридаги қарорда эса ҳеч қандай электрон рақамли имзо мавжуд эмас. Бу қарорнинг ваколатли шахс томонидан расмийлаштирилганлигини тасдиқлайдиган ягона ҳуқуқий асос бўлмиш электрон рақамли имзо мавжуд эмаслигини англатади. Қонунчиликка мувофиқ, тегишли тартибда тасдиқланмаган ҳужжат юридик кучга эга бўлмайди ва ноқонуний ҳисобланади.',
        
        'Иккинчидан, Вазирлар Маҳкамасининг 2018 йил 1 декабрдаги 975-сон қарорига мувофиқ қабул қилинган низомнинг 28-бандига кўра фото ва видео қайд этиш воситалари (радар, фоторадар) белгиланган тартибда сертификатланган бўлиши ҳамда ўз вақтида метрологик текширувдан ўтган бўлиши шарт.',
        
        'Бунда рухсат этилмаган, сертификатга эга бўлмаган, сертификатнинг амал қилиш муддати тугаган ёки белгиланган тартибда ички текширувдан ўтмаган махсус мосламалардан фойдаланиш қат\'иян тақиқланади.',
        
        'Аммо юқоридаги қарорда ҳеч қандай техник ҳужжат илова қилинмаган, бу эса ускунанинг қонунийлиги, аниқлик даражаси ва ишга яроқлилигини шубҳа остига қўяди.',
        
        'Ушбу низомнинг 29-бандига кўра "Транспорт воситасининг ҳайдовчиси махсус мослама қайд этган кўрсаткичдан норози бўлган тақдирда, ушбу ҳол бевосита жойнинг ўзида холисларнинг иштирокида расмийлаштирилади" деб белгиланган. Мазкур ишда ҳайдовчининг иштироки, эътирози ёки холислар мавжудлиги умуман қайд этилмаган.',
        
        'Учинчидан, Вазирлар Маҳкамасининг 975-сон қарорига мувофиқ қабул қилинган низомнинг 41-бандига кўра ЙПХ ходими йўлларда хизматни яширин олиб бориш тақиқланади. ЙПХ ходими ва техник воситалар яширин тарзда жойлаштирилмаслиги, ҳайдовчига уларнинг фаолияти очиқ ва ошкора бўлиши керак. Бироқ мазкур ишда на радар кўриниши ва на ЙПХ ходими фаолияти аниқ кўринган. Бу эса хизмат яширин олиб борилган бўлиши мумкинлигини кўрсатади ва шаффофлик тамойилига зид ҳисобланади.',
        
        'Тўртинчидан, Вазирлар Маҳкамасининг 2024 йил 25 мартдаги 148-сон қарори билан тасдиқланган низомнинг 2-иловасига кўра Махсус автоматлаштирилган фото ва видео қайд этиш тизимига доир техник талаблар белгилаб берилган. Ушбу ҳужжатнинг 6-параграф 22-бандига кўра "фото ва видео қайд этиш дастурий-техник воситалари орқали қайд этилган қоидабузарликлар тўғрисидаги қуйидаги маълумотлар «Маъмурий амалиёт» модулига тақдим этилиши шарт: қайд этилган транспорт воситасининг уч хил ҳолатдаги фотосурати; қоидабузарлик қайд этилган манзилнинг GPS координатаси; ҳолат акс этган 6 ± 10 сониядан иборат видеолавҳа ва унинг ҳаволаси.',
        
        'Бироқ, иш бўйича чиқарилган жарима солиш тўғрисидаги қарорда видеолавҳа мавжуд эмас. Бу эса қарорни техник ва ҳуқуқий талабларга жавоб бермаслигини кўрсатади.',
        
        'Бешинчидан, Вазирлар Маҳкамасининг 15.05.2025 йилдаги 318-сонли "Транспорт воситаларидан фойдаланишда аҳолига қўшимча қулайликлар яратиш чора-тадбирлари тўғрисида"ги қарори тўққизинчи бандига кўра "2025 йил 1 сентябрдан бошлаб ҳаракат тезлигини назорат қилишга мўлжалланган ҳар қандай турдаги стационар ва мобиль радарлар жойлашган ҳудудда ва унгача камида 250 метрдан кам ҳамда 500 метрдан кўп бўлмаган ҳудудда «Радар» йўл белгисининг ўрнатилиши мажбурий ҳисобланади. Бироқ, мазкур ҳолатда Радар йўл белгиси қўйилмаган эди. Бу қонунчилик талабларининг жиддий бузилиши ҳисобланиб қўлланилган маъмурий жазонинг ноқонунийлигини кўрсатади.',
        
        'Маъмурий жавобгарлик тўғрисидаги кодекси 321-моддаси иккинчи қисмига кўра Маъмурий ҳуқуқбузарлик тўғрисидаги иш юзасидан чиқарилган қарорни бекор қилиш учун маъмурий ҳуқуқбузарликлар тўғрисидаги ишларни юритиш қоидаларининг жиддий бузилиши; қўлланилган маъмурий жазонинг адолатсизлиги асос бўлиб ҳисобланади.',
        
        'Олтинчидан, Вазирлар Маҳкамасининг 2018 йил 1 декабрдаги 975-сон қарори билан тасдиқланган Низомнинг 34-бандида шундай белгиланган: «Кўчма фоторадар ва мобиль комплексларни қўллаш жойи ва вақти ДЙҲХХ саф бўлими (бўлинмаси) бошлиғи томонидан тасдиқланган дислокацияга мувофиқ белгиланади».',
        
        'Агар радар хусусий тадбиркорларга тегишли бўлса: Вазирлар Маҳкамасининг 2021 йил 21 апрелдаги 232-сон қарорига мувофиқ, тадбиркорлар радарларни фақат «E-auksion» савдо майдончаси орқали ютиб олинган, ИИВ томонидан белгилаб берилган аниқ дислокация (координата) нуқтасигагина ўрнатишга ҳақли. Уни исталган бошқа жойга кўчириб юриши мутлақо ноқонуний.',
        
        'Мазкур ҳолатда иш юритиш қоидалари жиддий равишда бузилган бўлиб, бу МЖтКнинг 321-моддаси 2-қисмида назарда тутилган асосдир. Унга кўра, иш юритишда жиддий процессуал ҳатоларга йўл қўйилган бўлса, бунда чиқарилган қарор қонуний кучга эга бўлмайди ва бекор қилиниши керак. Юқорида санаб ўтилган барча қонунбузилишлар биргаликда ушбу модда доирасида иш юритишни тугатиш учун асос бўлади.',
        
        'Бундан ташқари, Ўзбекистон Республикаси Конституциясининг 20-моддасида "Давлат органлари томонидан инсонга нисбатан қўлланиладиган ҳуқуқий таъсир чоралари мутаносиблик принципига асосланиши ва қонунларда назарда тутилган мақсадларга эришиш учун етарли бўлиши керак. Инсон билан давлат органларининг ўзаро муносабатларида юзага келадиган қонунчиликдаги барча зиддиятлар ва ноаниқликлар инсон фойдасига талқин этилади" деб қатъий белгилаб қўйилган.',
        
        'Юқоридагиларга кўра, Конституциянинг 20-моддаси ва МЖтКнинг 171, 271, 283, 309¹, 321-моддаларига асосан,',
    ]
    
    for text in body_paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        # Paragraflar orasidagi bo'shliq
        p_format = p.paragraph_format
        p_format.space_after = Pt(6)
        if text.startswith('\t'):
            p_format.first_line_indent = Cm(1.25)
    
    # ===== SO'RAYMAN =====
    doc.add_paragraph()  # Bo'sh qator
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('С Ў Р А Й М А Н :')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('{{ORGAN_NOMI}}нинг {{QAROR_SANASI}} йилдаги {{QAROR_RAQAMI}}-сонли қарорини бекор қилишингизни.')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # ===== ILOVA =====
    p = doc.add_paragraph()
    run = p.add_run('Илова:')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    
    ilova_items = [
        'Жарима солиш тўғрисидаги қарор;',
        'Ҳайдовчилик гувоҳномаси;',
        'Автомототранспорт воситаси рўйхатдан ўтказилганлиги тўғрисидаги гувоҳнома;',
        'Шахсни тасдиқловчи ҳужжат (паспорт/ID карта).'
    ]
    
    for item in ilova_items:
        p = doc.add_paragraph()
        run = p.add_run(f'- {item}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
    
    # ===== IMZO =====
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('{{SANA}}\t\t       \t\t(имзо)\t\t\t    {{FISH}}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    # Saqlash
    output_path = r'd:\Apps\coding\UchOyoqBot\templates\ariza_krill.docx'
    doc.save(output_path)
    print(f"Shablon yaratildi: {output_path}")


def create_template_latin():
    """Lotin alifbosidagi ariza shablonini yaratish."""
    doc = Document()
    
    # Sahifa sozlamalari
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    
    # Default shrift
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    # ===== HEADER =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Jinoyat ishlari bo'yicha {{TUMAN_NOMI}} sudiga")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Arizachi: {{FISH}}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Yashash manzil: {{MANZIL}}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Telefon raqam: {{TEL}}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # ===== TITLE =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('A R I Z A')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("({{QAROR_RAQAMI}}-sonli jarimani bekor qilish to'g'risida)")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # ===== BODY =====
    body_paragraphs = [
        "Men, {{FISH}}, {{AVTOMOBIL}} rusumli avtomobil davlat raqami {{DAVLAT_RAQAMI}} avtomobilga {{ORGAN_NOMI}} tomonidan MJtKning {{MODDA}} bilan huquqbuzarlik sodir etganlikda aybdor deb topildim va ushbu huquqbuzarlik yuzasidan {{QAROR_SANASI}} kuni {{QAROR_RAQAMI}}-sonli qaror rasmiylashtirildi.",
        
        "Ushbu {{QAROR_RAQAMI}}-sonli jarima solish to'g'risidagi qarorni O'zbekiston Respublikasining amaldagi qonunlari juda QO'POL ravishda buzilgan deb hisoblaymiz. Qaror quyidagi qonuniy asoslar buzilgan holda rasmiylashtirilgan:",
        
        "Birinchi dalil: Vazirlar Mahkamasining 2018-yil 1-dekabrdagi 975-sonli qarori:",

        "Nizomning 6-bobi 1-paragrafi buzilgan",
        
        "975-sonli Hukumat Qarori asosida tasdiqlangan Nizomning 6-bobi 1-paragrafida aniq ko'rsatilganidek, ko'chma fotoradarlar, mobil foto va video qayd etish komplekslari orqali aniqlangan qoidabuzarlik holatlarida to'g'ridan-to'g'ri QAROR emas, faqatgina BAYONNOMA rasmiylashtirilishi kerak. Bu me'yor yo'l harakati qoidabuzarligi bo'yicha haydovchining ishtirokida holis va shaffof ish yuritishni ta'minlashga qaratilgan. Ammo amalda menga nisbatan bayonnoma tuzilmasdan, bevosita qaror chiqarilgan, bu esa mazkur normativ-huquqiy hujjat talablari qo'pol ravishda buzilganini anglatadi.",

        "Nizomning 28-bandi buzilgan",
        
        "Nizomning 28-bandiga binoan, sertifikatsiyadan o'tmagan, yoki amal qilish muddati o'tgan texnik vositalardan foydalanish qat'iyan man etiladi. Ayniqsa fotoradarlar yoki boshqa nazorat komplekslari ishlatilgan hollarda, ularning amaldagi sertifikati, metrologik tekshiruvdan o'tganlik guvohnomasi, kalibrovka hujjatlari va texnik pasportlari qaror ilovasida mavjud bo'lishi shart. Ammo {{QAROR_RAQAMI}}-sonli qarorga hech qanday texnik hujjat ilova qilinmagan, bu esa bu uskunaning qonuniyligini, aniqlik darajasini va ishga yaroqliligini shubha ostiga qo'yadi.",

        "Nizomning 29-bandi buzilgan",
        
        "Nizomning 29-bandida qayd etilishicha, agar haydovchi radar ko'rsatkichidan norozi bo'lsa, bu holat albatta holislar ishtirokida rasmiylashtirilishi lozim. Bu normalarning amaliy ahamiyati – haydovchining e'tirozini qonuniy tartibda qayd qilish va uning fikrini inkor qilinmas huquqiy dalil sifatida e'tirof etishdir. Mazkur ishda esa haydovchining ishtiroki, e'tirozi yoki holislar mavjudligi umuman qayd etilmagan, bu esa shaxsning isbotlash huquqini buzgan holda, ishni bir tomonlama ko'rib chiqishga olib kelgan.",

        "Nizomning 37-bobi 2-qismi buzilgan",
        
        "Nizomning 37-bobi 2-qismiga muvofiq, foto va video fiksatsiyaga asoslangan qarorlar majburiy ravishda elektron raqamli imzo (ERI) bilan tasdiqlangan bo'lishi shart. Bu talab hujjatning rasmiyligini, qonuniy kuchga egaligini va sud hamda boshqa idoralarda dalil sifatida qabul qilinishini ta'minlaydi. {{QAROR_RAQAMI}}-sonli qarorda esa hech qanday elektron imzo mavjud emas, bu esa hujjatni yuridik kuchdan mahrum qiladi va uni qonuniy deb hisoblashga to'sqinlik qiladi.",

        "Nizomning 41-bandi buzilgan",
        
        "Nizomning 41-bandida ochiq-oydin belgilangan: yo'l-patrul xizmati faoliyati oshkora olib borilishi shart. YPX xodimi va texnik vositalar yashirin tarzda joylashtirilmasligi, haydovchiga ularning mavjudligi ayon bo'lishi kerak. Biroq mazkur ishda na radar ko'rinishi, na YPX xodimi faoliyati aniq sezilgan. Bu esa xizmat yashirin olib borilgan bo'lishi mumkinligini ko'rsatadi va shaffoflik tamoyiliga zid hisoblanadi.",

        "MJtKning 321-moddasi 2-qismi buzilgan",
        
        "Mazkur holatda ish yuritish qoidalari jiddiy ravishda buzilgan bo'lib, bu esa MJtKning 321-moddasi 2-qismida ko'zda tutilgan asosdir. Unga ko'ra, agar ish yuritishda jiddiy protsessual xatolarga yo'l qo'yilgan bo'lsa, bunda chiqarilgan qaror qonuniy kuchga ega bo'lmaydi va bekor qilinishi kerak. Yuqorida sanab o'tilgan barcha qonunbuzilishlar birgalikda ushbu modda doirasida ish yuritishni tugatish uchun asos bo'ladi.",
        
        "Bundan tashqari, O'zbekiston Respublikasi Konstitutsiyasining 20-moddasida quyidagilar belgilangan:",
        
        '"Davlat organlari tomonidan insonga nisbatan qo\'llaniladigan huquqiy ta\'sir choralari mutanosiblik prinsipiga asoslanishi va qonunlarda nazarda tutilgan maqsadlarga erishish uchun yetarli bo\'lishi kerak. Inson bilan davlat organlarining o\'zaro munosabatlarida yuzaga keladigan qonunchilikdagi barcha ziddiyatlar va noaniqliklar inson foydasiga talqin etiladi."',
        
        "Mazkur holatda esa, nafaqat mutanosiblik prinsipi buzilgan, balki qaror chiqarishda mavjud ziddiyatlar fuqaro – ya'ni meni foydamga talqin etilmagan.",
        
        "Ikkinchi dalil: Shu bilan birga, 148-sonli qarorning 6 bobi 22-bandida belgilanishicha, foto va video qayd etish vositalari orqali aniqlangan huquqbuzarlik holatlari \"Ma'muriy amaliyot\" axborot tizimiga quyidagi majburiy axborotlar bilan taqdim etilishi shart deb belgilangan: uch xil rakursdagi fotosuratlar, harakat tezligi, aniq manzil, sana va 6–10 soniyalik videolavha yoki videoga havola.",
        
        "Mazkur qarorda yuqoridagi majburiy dalillarning eng asosiy elementi – videoyozuv mavjud emasligi sababli qaror texnik va huquqiy talablarga javob bermaydi.",
        
        "Uchinchi dalil: O'zbekiston Respublikasi Prezidentining 2018-yil 30-oktabrdagi PQ-3989-sonli qarorining 4-bandi quyidagilarga ruxsat beradi:\n\"Investor bilan hamkorlikda Qoraqalpog'iston Respublikasi va viloyatlarda ko'k va qizil rangli yalt-yalt etuvchi chiroq-mayoqcha, maxsus tovushli ishoralari hamda avtomatlashtirilgan jamoat tartibini nazorat qilish va yo'l harakati qoidabuzarliklarini aniqlash moslamalari o'rnatilgan PATRUL avtomobillarini amaliyotga joriy qilish.\"",
        
        "Mazkur xizmatni aynan patrul xodimi olib borgan-bormaganligi ham shubha ostida qolmoqda. Bu holat xizmatning qonuniyligi va javobgarlikka tortishning asoslanganligiga jiddiy savollar tug'diradi.",
        
        "Qolaversa, haydovchi na patrul xizmat transport vositasini, na unda faoliyat yuritayotgan patrul xodimini umuman ko'rmagan. Bu esa xizmatning qanday sharoitda amalga oshirilgani borasida jiddiy shubhalar uyg'otadi.",
        
        "Vazirlar Mahkamasining 975-sonli qarori 8-bobi 41-bandiga ko'ra, yo'l-patrul xizmatining faoliyati yashirin tarzda emas, oshkora shaklda olib borilishi shart. Ushbu talabga rioya etilmaganligi xizmatning o'zini ham, uning asosida chiqarilgan qarorni ham qonuniylikdan mahrum qiladi.",
        
        "Nizomning 32-bandi 2-qismiga muvofiq, quyidagicha qat'iy belgilangan:\n\"Mazkur talabga rioya etilmagan holda rasmiylashtirilgan bayonnomalar yuridik kuchga ega bo'lmaydi va huquqiy oqibatlar keltirib chiqarmaydi.\"",
        
        "Shu asosda, agar bayonnoma qonun bilan belgilangan talablarga rioya qilinmasdan rasmiylashtirilgan bo'lsa, u holda mazkur bayonnoma (yoki unga asoslangan qaror) yuridik kuchga ega emas va unga asosan fuqaroga nisbatan hech qanday huquqiy ta'sir chorasi qo'llanishi mumkin emas.",
        
        "O'zbekiston Respublikasi Ma'muriy javobgarlik to'g'risidagi kodeksining 321-moddasi 2-qismida quyidagicha belgilangan:\n\"Ma'muriy huquqbuzarliklar to'g'risidagi ishlarni yuritish qoidalarining jiddiy buzilishi – ma'muriy huquqbuzarlik to'g'risidagi ish yuzasidan chiqarilgan qarorni bekor qilishga asos bo'ladi.\"",
        
        "Yuqoridagi qonuniy faktlarga asosan, O'zbekiston Respublikasi MJtKning {{MODDA}}siga asosan {{JARIMA_SUMMASI}} so'm miqdorida jarima belgilangan, {{QAROR_SANASI}} sanadagi {{QAROR_RAQAMI}}-sonli qarorni MJtKning 271-moddasi 1-bandiga asosan bekor qilishingizni so'rayman.",
    ]
    
    for text in body_paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        p_format = p.paragraph_format
        p_format.space_after = Pt(6)
    
    doc.add_paragraph()
    
    # ===== ILOVA =====
    p = doc.add_paragraph()
    run = p.add_run('Ilova:')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    
    ilova_items = [
        "Haydovchilik guvohnomasi;",
        "Avtomototransport vositasi ro'yxatdan o'tkazilganligi to'g'risidagi guvohnoma;",
        "Shaxsni tasdiqlovchi hujjat (pasport/ID karta);",
        f"Jarima qarori nusxasi."
    ]
    
    for item in ilova_items:
        p = doc.add_paragraph()
        run = p.add_run(f'- {item}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('IMZO:   _________________      {{FISH}}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    # Saqlash
    output_path = r'd:\Apps\coding\UchOyoqBot\templates\ariza_lotin.docx'
    doc.save(output_path)
    print(f"Shablon yaratildi: {output_path}")


if __name__ == "__main__":
    create_template()
    create_template_latin()
    print("Barcha shablonlar tayyor!")
