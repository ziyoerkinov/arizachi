"""
Doc Generator — DOCX shablonni to'ldirish va PDF ga o'girish moduli.

1. python-docx orqali shablon ochib, {{placeholder}} larni haqiqiy qiymatlar bilan almashtiradi
2. docx2pdf (MS Word) orqali PDF ga o'giradi
"""

import os
import re
from datetime import datetime
from docx import Document
from docx2pdf import convert

import config


def fill_template(template_path: str, data: dict, output_docx_path: str) -> bool:
    """
    DOCX shablondagi {{placeholder}} larni haqiqiy qiymatlar bilan almashtirish.
    
    Args:
        template_path: Shablon fayl yo'li
        data: Ajratilgan ma'lumotlar (extractor dan)
        output_docx_path: Natija DOCX fayl yo'li
        
    Returns:
        bool: Muvaffaqiyat holati
    """
    try:
        doc = Document(template_path)
        
        # Bugungi sana
        today = datetime.now().strftime("%d.%m.%Y")
        
        # Placeholder → qiymat mapping
        replacements = {
            "{{FISH}}": data.get("fish", "____________________"),
            "{{MANZIL}}": data.get("manzil", "____________________"),
            "{{TEL}}": data.get("tel", "____________________"),
            "{{QAROR_RAQAMI}}": data.get("qaror_raqami", "____________"),
            "{{QAROR_SANASI}}": data.get("qaror_sanasi", "____.____._______"),
            "{{MODDA}}": data.get("modda", "128-modda"),
            "{{JARIMA_SUMMASI}}": _format_summa(data.get("jarima_summasi", "")),
            "{{ORGAN_NOMI}}": data.get("organ_nomi", "____________________"),
            "{{TUMAN_NOMI}}": data.get("tuman_nomi", "________________"),
            "{{AVTOMOBIL}}": data.get("avtomobil", "____________________"),
            "{{DAVLAT_RAQAMI}}": data.get("davlat_raqami", "________"),
            "{{SANA}}": today,
        }
        
        # Barcha paragraflarda almashtirish
        for paragraph in doc.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        
        # Jadval ichida ham (agar bo'lsa)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_in_paragraph(paragraph, replacements)
        
        # Header va footer da ham
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)
        
        doc.save(output_docx_path)
        print(f"[DocGen] DOCX saqlandi: {output_docx_path}")
        return True
        
    except Exception as e:
        print(f"[DocGen] Shablon to'ldirish xatolik: {e}")
        import traceback
        traceback.print_exc()
        return False


def _replace_in_paragraph(paragraph, replacements: dict):
    """
    Paragraf ichidagi run'larda placeholder almashtirishni amalga oshirish.
    
    Ba'zan Word bir placeholder ni bir nechta run ga bo'lib qo'yadi.
    Shu sababli avval barcha run'larni birlashtirib, keyin almashtiramiz.
    """
    # Paragrafning to'liq matnini olish
    full_text = paragraph.text
    
    # Agar hech qanday placeholder bo'lmasa, o'tkazib yuborish
    if "{{" not in full_text:
        return
    
    # Almashtirishni amalga oshirish
    new_text = full_text
    for placeholder, value in replacements.items():
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, value)
    
    # Agar o'zgarish bo'lmasa, o'tkazib yuborish
    if new_text == full_text:
        return
    
    # Run'larni yangilash — birinchi run ga to'liq matnni qo'yish,
    # qolganlarni tozalash (formatlashni saqlash uchun)
    if paragraph.runs:
        # Birinchi run formatini saqlash
        first_run = paragraph.runs[0]
        first_run.text = new_text
        
        # Qolgan run'larni tozalash
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = new_text


def _format_summa(summa_str: str) -> str:
    """Jarima summasini formatlash (masalan: 412000 -> 412 000)."""
    if not summa_str:
        return "____________"
    
    # Faqat raqamlarni olish
    digits = re.sub(r'\D', '', summa_str)
    
    if not digits:
        return summa_str
    
    # Minglik ajratgich bilan formatlash
    try:
        number = int(digits)
        return f"{number:,}".replace(",", " ")
    except ValueError:
        return summa_str


def convert_to_pdf(docx_path: str, output_dir: str = None) -> str:
    """
    DOCX faylni PDF ga o'girish (MS Word orqali).
    
    Args:
        docx_path: DOCX fayl yo'li
        output_dir: PDF saqlanadigan papka (None bo'lsa, DOCX bilan bir papkada)
        
    Returns:
        str: Yaratilgan PDF fayl yo'li yoki None
    """
    try:
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
            pdf_filename = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
            pdf_path = os.path.join(output_dir, pdf_filename)
        else:
            pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
        
        print(f"[DocGen] PDF ga o'girish: {docx_path} -> {pdf_path}")
        convert(docx_path, pdf_path)
        print(f"[DocGen] PDF tayyor: {pdf_path}")
        return pdf_path
        
    except Exception as e:
        print(f"[DocGen] PDF konvertatsiya xatolik: {e}")
        import traceback
        traceback.print_exc()
        return None


def generate_document(data: dict, template_type: str = "krill") -> str:
    """
    To'liq jarayon — shablon to'ldirish va PDF yaratish.
    
    Args:
        data: Ajratilgan ma'lumotlar dict
        template_type: "krill" yoki "lotin"
        
    Returns:
        str: Yaratilgan PDF fayl yo'li yoki None
    """
    # Shablon yo'li
    template_filename = f"ariza_{template_type}.docx"
    template_path = os.path.join(config.TEMPLATES_DIR, template_filename)
    
    if not os.path.exists(template_path):
        print(f"[DocGen] Shablon topilmadi: {template_path}")
        return None
    
    # Chiqish fayl nomlari
    qaror_raqam = data.get("qaror_raqami", "nomalum")
    # Fayl nomidan noto'g'ri belgilarni tozalash
    safe_name = re.sub(r'[<>:"/\\|?*]', '_', qaror_raqam)
    
    output_docx = os.path.join(config.OUTPUT_DIR, f"Ariza_{safe_name}.docx")
    
    # Papkani yaratish
    os.makedirs(config.OUTPUT_DIR, exist_ok=True)
    
    # 1-qadam: Shablon to'ldirish
    if not fill_template(template_path, data, output_docx):
        return None
    
    # 2-qadam: PDF ga o'girish
    pdf_path = convert_to_pdf(output_docx, config.OUTPUT_DIR)
    
    # Vaqtinchalik DOCX ni o'chirish (ixtiyoriy)
    # os.remove(output_docx)
    
    return pdf_path


if __name__ == "__main__":
    # Test
    test_data = {
        "fish": "Obilov Otamurod Boymurodovich",
        "manzil": "Qashqadaryo viloyati, Qamashi tumani, Chim QFY",
        "tel": "+998 90 123 45 67",
        "qaror_raqami": "SD25700146635",
        "qaror_sanasi": "28.10.2025",
        "modda": "128-modda 3-qismi",
        "jarima_summasi": "412000",
        "organ_nomi": "Qashqadaryo viloyati MAB inspektori S.Shodixonov",
        "tuman_nomi": "Qamashi tumani",
        "avtomobil": "Nexia",
        "davlat_raqami": "85A123BC"
    }
    
    # Krill shablon test
    result = generate_document(test_data, "krill")
    if result:
        print(f"\nKrill PDF tayyor: {result}")
    
    # Lotin shablon test
    result = generate_document(test_data, "lotin")
    if result:
        print(f"\nLotin PDF tayyor: {result}")
